#!/usr/bin/env python3
"""Генерация документации Харбор в формате DOCX для защиты проекта."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Харбор-документация.docx"


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(80, 80, 80)


def h1(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = "Times New Roman"


def h2(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = "Times New Roman"


def h3(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = "Times New Roman"


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Times New Roman"


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Times New Roman"


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Number")
    for run in p.runs:
        run.font.name = "Times New Roman"


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1)


def build() -> None:
    doc = Document()
    set_style(doc)

    # Титул
    add_title(doc, "Харбор")
    add_subtitle(doc, "Веб-приложение для поиска и бронирования жилья")
    add_subtitle(doc, "Документация для защиты проекта")
    doc.add_paragraph()

    # 1. О проекте
    h1(doc, "1. О проекте")
    para(doc,
         "Харбор — одностраничное веб-приложение (SPA), аналог Airbnb. "
         "Пользователь может искать жильё по городу и датам, просматривать карточки объектов, "
         "добавлять в избранное, бронировать, управлять профилем и поездками. "
         "Интерфейс поддерживает русский и английский языки.")
    para(doc, "Основные возможности:")
    bullet(doc, "Поиск жилья по городу, датам заезда/выезда и количеству гостей")
    bullet(doc, "Фильтрация результатов (цена, тип жилья, удобства, доступность по датам)")
    bullet(doc, "Страница объекта с галереей, описанием, отзывами, картой района")
    bullet(doc, "Бронирование с выбором дат, гостей, способа оплаты")
    bullet(doc, "Регистрация, вход, профиль, избранное, история поездок")
    bullet(doc, "Информационные страницы, новости, идеи для путешествий")
    bullet(doc, "Адаптивная вёрстка: мобильные (≤767px), планшет (≤1023px), десктоп (≥1024px)")

    # 2. Технологический стек
    h1(doc, "2. Технологический стек")
    h2(doc, "2.1. Frontend")
    bullet(doc, "React 19 + TypeScript — UI-компоненты и бизнес-логика")
    bullet(doc, "Vite 8 — сборка, dev-сервер, HMR")
    bullet(doc, "react-router-dom 7 — маршрутизация (SPA)")
    bullet(doc, "SCSS Modules — изолированные стили компонентов")
    bullet(doc, "date-fns — работа с датами в календаре и бронировании")
    bullet(doc, "axios — HTTP-запросы к backend API")
    bullet(doc, "lucide-react — иконки")

    h2(doc, "2.2. Backend")
    bullet(doc, "Express (Node.js) — REST API на порту 5001")
    bullet(doc, "PostgreSQL 16 — база данных (Docker, порт 5433)")
    bullet(doc, "JWT + bcrypt — аутентификация и хеширование паролей")
    bullet(doc, "json-server (опционально) — mock API для каталога объектов на порту 3001")

    h2(doc, "2.3. Данные объектов")
    bullet(doc, "src/db.json — статический каталог из 26 объектов недвижимости")
    bullet(doc, "Каждый объект содержит: название, город, цену, фото, описание, удобства, хоста, занятые даты (bookedDates)")

    # 3. Структура проекта
    h1(doc, "3. Структура проекта")
    code_block(doc, """харбор/
├── public/images/properties/   # Фото объектов (1–26)
├── src/
│   ├── App.tsx                 # Корневой компонент, маршруты
│   ├── main.tsx                # Точка входа React
│   ├── db.json                 # Каталог объектов
│   ├── assets/images/          # Локальные изображения (trip-ideas, home)
│   ├── components/             # UI-компоненты и блоки
│   │   ├── layout/             # Header, Footer, MainLayout, PageLayout
│   │   ├── ui/                 # Button, Input, DatePicker, PropertyCard…
│   │   ├── blocks/             # PropertiesSection, PopularDestinations…
│   │   └── features/           # MobileSearchFlow, PersonalInfoTab…
│   ├── pages/                  # Страницы приложения
│   ├── context/                # React Context (Auth, Search, Favorites, Language)
│   ├── i18n/                   # Переводы ru/en
│   ├── utils/                  # Вспомогательные функции
│   ├── hooks/                  # Кастомные хуки
│   ├── types/                  # TypeScript-типы
│   ├── styles/                 # variables.scss, mixins.scss
│   └── harbor-backend/         # Express API + PostgreSQL
├── scripts/                    # Скрипты (retina-изображения, deploy-check)
├── docker-compose.yml          # PostgreSQL контейнер
└── vite.config.ts              # Конфиг Vite + прокси /api → backend""")

    # 4. Архитектура
    h1(doc, "4. Архитектура приложения")
    h2(doc, "4.1. Точка входа и провайдеры")
    para(doc,
         "main.tsx монтирует React-приложение в DOM и оборачивает его в BrowserRouter. "
         "App.tsx включает цепочку провайдеров контекста:")
    bullet(doc, "LanguageProvider — текущий язык (ru/en), словарь переводов")
    bullet(doc, "AuthProvider — авторизация, JWT-токен, данные пользователя, аватар")
    bullet(doc, "SearchProvider — параметры поиска (город, даты, гости), URL-построение")
    bullet(doc, "FavoritesProvider — синхронизация избранного с backend")
    bullet(doc, "AppShell — общая оболочка (скролл, модалки)")

    h2(doc, "4.2. Маршрутизация (App.tsx)")
    para(doc, "Приложение использует react-router-dom. Маршруты сгруппированы по layout:")
    bullet(doc, "/ — главная (HomePage)")
    bullet(doc, "/search?city=… — результаты поиска")
    bullet(doc, "/property/:id — страница объекта")
    bullet(doc, "/property/:id/book — оформление бронирования")
    bullet(doc, "/favorites — избранное")
    bullet(doc, "/profile — профиль пользователя")
    bullet(doc, "/ideas/:slug — статья «Идеи для путешествий»")
    bullet(doc, "/login, /register — авторизация")
    bullet(doc, "/about, /news, /privacy, /terms, /contact и др. — информационные страницы")

    para(doc,
         "MainLayout управляет отображением Header (с поисковой строкой или без), Footer и отступами. "
         "На странице бронирования Footer скрыт.")

    h2(doc, "4.3. Layout-компоненты")
    bullet(doc, "Header — логотип, SearchBar (desktop), UserMenu, переключатель языка")
    bullet(doc, "SearchBar — поля «Куда», «Прибытие», «Выезд», «Гости»")
    bullet(doc, "MobileSearchFlow — пошаговый поиск на мобильных")
    bullet(doc, "MainLayout — обёртка страниц с Header/Footer")
    bullet(doc, "PageLayout — контейнер с max-width и отступами для контентных страниц")
    bullet(doc, "Footer — ссылки на разделы сайта")

    # 5. Ключевые модули
    h1(doc, "5. Ключевые модули и как они работают")

    h2(doc, "5.1. Каталог объектов (db.json)")
    para(doc,
         "Файл src/db.json содержит массив properties из 26 объектов. Каждый объект имеет поля: "
         "id, title, location, city, pricePerNight, rating, reviewsCount, imageUrl, images[], "
         "description, amenities[], host{}, locationDetails{}, bookedDates[], propertyType, amenityIds.")
    para(doc,
         "Данные загружаются напрямую из db.json (import) или через json-server на localhost:3001. "
         "При недоступности mock-сервера используется fallback на локальный db.json.")

    h2(doc, "5.2. Поиск и фильтрация")
    para(doc, "SearchContext хранит состояние поиска и синхронизирует его с URL и localStorage.")
    para(doc, "SearchResultsPage.tsx:")
    bullet(doc, "Читает параметры city, checkIn, checkOut, guests из URL")
    bullet(doc, "Загружает все объекты из db.json")
    bullet(doc, "Применяет applySearchFilters (цена, тип, удобства, рейтинг)")
    bullet(doc, "filterAvailableForDates — исключает объекты с пересечением bookedDates")
    bullet(doc, "FilterSidebar — боковая панель фильтров на desktop")

    h2(doc, "5.3. Календарь и занятые даты")
    para(doc,
         "DatePicker — универсальный компонент выбора дат. Поддерживает режимы: inline, modal, scroll (mobile). "
         "Занятые даты (bookedDates из db.json) отображаются с диагональным зачёркиванием и недоступны для выбора.")
    para(doc, "Утилиты:")
    bullet(doc, "propertyBookings.ts — проверка пересечения дат, получение занятых диапазонов")
    bullet(doc, "searchFilters.ts — filterAvailableForDates для поиска")

    h2(doc, "5.4. Страница объекта (PropertyDetails)")
    para(doc, "PropertyDetailsPage загружает объект по id из db.json и отображает:")
    bullet(doc, "ImageGallery — сетка фото + модальная галерея")
    bullet(doc, "Описание, удобства, информация о хосте")
    bullet(doc, "LocationSection — адрес, транспорт, инфраструктура")
    bullet(doc, "ReviewsSection — отзывы")
    bullet(doc, "ReservationCard — sidebar с ценой, календарём, кнопкой «Забронировать»")

    h2(doc, "5.5. Бронирование")
    para(doc, "Поток: PropertyDetails → ReservationCard → /property/:id/book (BookingRequestPage)")
    bullet(doc, "BookingRequestSidebar — превью объекта, итоговая цена, кнопка подтверждения")
    bullet(doc, "BookingEditModal — редактирование дат, гостей, правил отмены (mode='info')")
    bullet(doc, "При подтверждении — POST /api/bookings с JWT-токеном")
    bullet(doc, "Backend проверяет конфликт дат (hasBookingConflict) и сохраняет в PostgreSQL")

    h2(doc, "5.6. Аутентификация")
    para(doc,
         "AuthContext управляет сессией. Токен и данные пользователя хранятся в localStorage/sessionStorage "
         "(authStorage.ts). LoginPage и RegisterPage отправляют запросы на POST /api/login и POST /api/register.")
    para(doc, "Backend: bcrypt для паролей, JWT (24h) для токенов, middleware authenticateToken для защищённых маршрутов.")

    h2(doc, "5.7. Избранное")
    para(doc,
         "FavoritesContext синхронизирует список propertyId с backend: "
         "GET/POST/DELETE /api/favorites. На клиенте также используется localStorage как кэш.")

    h2(doc, "5.8. Профиль и поездки")
    para(doc, "ProfilePage содержит вкладки:")
    bullet(doc, "PersonalInfoTab — имя, email, телефон, смена пароля, аватар")
    bullet(doc, "TripsTab — список бронирований из GET /api/bookings")
    bullet(doc, "FavoritesTab — избранные объекты")

    h2(doc, "5.9. Интернационализация (i18n)")
    para(doc,
         "LanguageContext + useTranslation(). Переводы в src/i18n/translations/ru/ и en/. "
         "localizeProperties() подставляет локализованные названия городов и типов жилья.")

    h2(doc, "5.10. Компонент ResponsiveImage")
    para(doc,
         "Для локальных .webp-файлов автоматически добавляет srcset с @2x версией "
         "(buildRetinaSrcSet). Используется в PropertyCard, SearchResultCard, ImageGallery, BookingConfirmModal.")

    # 6. Backend API
    h1(doc, "6. Backend API (Express + PostgreSQL)")
    para(doc, "Сервер: src/harbor-backend/index.ts, порт 5001.")
    para(doc, "Vite проксирует /api/* → http://localhost:5001/* (vite.config.ts).")
    para(doc, "Таблицы PostgreSQL:")
    bullet(doc, "users — пользователи (id, first_name, last_name, email, password_hash, phone)")
    bullet(doc, "favorites — избранное (user_id, property_id)")
    bullet(doc, "bookings — бронирования (даты, гости, цена, статус, snapshot объекта)")
    bullet(doc, "host_applications — заявки на размещение жилья")

    para(doc, "Основные эндпоинты:")
    bullet(doc, "GET /health — проверка БД")
    bullet(doc, "POST /register, POST /login — регистрация и вход")
    bullet(doc, "GET/PATCH /users/me, PUT /users/me/password — профиль")
    bullet(doc, "GET/POST/DELETE /favorites — избранное")
    bullet(doc, "GET/POST/DELETE /bookings — бронирования")
    bullet(doc, "POST /host-applications — заявка хозяина")

    # 7. Запуск
    h1(doc, "7. Запуск проекта")
    numbered(doc, "Установить зависимости: npm install && npm install --prefix src/harbor-backend")
    numbered(doc, "Запустить PostgreSQL: npm run db:up (docker-compose)")
    numbered(doc, "Запустить всё: npm start (backend + frontend)")
    numbered(doc, "Или по отдельности: npm run backend и npm run dev")
    numbered(doc, "Frontend: http://localhost:5173")
    numbered(doc, "Backend API: http://localhost:5001")
    numbered(doc, "Сборка production: npm run build")

    # 8. Стили
    h1(doc, "8. Система стилей")
    bullet(doc, "styles/variables.scss — цвета, отступы, breakpoints, z-index")
    bullet(doc, "styles/mixins.scss — media queries, типографика")
    bullet(doc, "Каждый компонент имеет свой *.module.scss")
    bullet(doc, "index.css — глобальные CSS-переменные и reset")
    bullet(doc, "Breakpoints: mobile ≤767px, tablet-down ≤1023px, desktop ≥1024px")

    # 9. ЗАМЕНА ФОТОГРАФИЙ
    h1(doc, "9. Как вручную заменить все фотографии в проекте")
    para(doc,
         "В проекте используются три типа изображений: локальные файлы (.webp), "
         "URL с Unsplash (внешние ссылки) и аватары пользователей (localStorage). "
         "Ниже — полная инструкция по каждой категории.")

    h2(doc, "9.1. Общие правила")
    bullet(doc, "Предпочтительный формат локальных файлов — WebP")
    bullet(doc, "Для Retina-экранов нужны две версии: card.webp (1x) и card@2x.webp (2x)")
    bullet(doc, "После замены 1x-файла запустите: npm run images:retina — скрипт автоматически создаст @2x")
    bullet(doc, "Скрипт scripts/generate-retina-images.mjs обрабатывает public/images/properties/ и src/assets/images/")
    bullet(doc, "Компонент ResponsiveImage автоматически подставляет srcset для .webp файлов")

    h2(doc, "9.2. Фото объектов недвижимости (26 штук)")
    para(doc, "Это основные карточки жилья. Каждый объект имеет папку с id от 1 до 26.")
    para(doc, "Путь к файлам:")
    code_block(doc, "public/images/properties/{id}/card.webp\npublic/images/properties/{id}/card@2x.webp")
    para(doc, "Пошаговая замена карточки объекта (например, объект id=5):")
    numbered(doc, "Подготовьте новое фото. Рекомендуемый размер: ~800×600 px (1x)")
    numbered(doc, "Конвертируйте в WebP (можно через squoosh.app, Photoshop или: cwebp input.jpg -o card.webp)")
    numbered(doc, "Замените файл: public/images/properties/5/card.webp")
    numbered(doc, "Запустите в терминале: npm run images:retina")
    numbered(doc, "Проверьте, что появился/обновился public/images/properties/5/card@2x.webp")
    numbered(doc, "Откройте src/db.json, найдите объект с \"id\": 5")
    numbered(doc, "Убедитесь, что поля imageUrl и images[0] указывают на /images/properties/5/card.webp")
    numbered(doc, "Перезапустите dev-сервер (npm run dev) и проверьте в браузере")

    para(doc, "Повторите для всех 26 объектов (папки 1–26 в public/images/properties/).")

    h3(doc, "Дополнительные фото в галерее объекта (images[] в db.json)")
    para(doc,
         "У объектов с id 1–10 в массиве images[] помимо локального card.webp есть 3 URL с Unsplash. "
         "Они показываются в галерее на странице объекта и в карусели на mobile.")
    para(doc, "Чтобы заменить их на свои локальные фото:")
    numbered(doc, "Положите новые файлы в папку объекта, например: public/images/properties/1/photo2.webp, photo3.webp, photo4.webp")
    numbered(doc, "Запустите npm run images:retina")
    numbered(doc, "В db.json замените URL Unsplash на локальные пути:")
    code_block(doc, '"images": [\n  "/images/properties/1/card.webp",\n  "/images/properties/1/photo2.webp",\n  "/images/properties/1/photo3.webp",\n  "/images/properties/1/photo4.webp"\n]')
    para(doc, "У объектов 11–26 в images[] только один локальный файл — card.webp.")

    h3(doc, "Аватары хостов (host.avatar в db.json)")
    para(doc,
         "У всех 26 объектов поле host.avatar содержит URL Unsplash. "
         "Чтобы заменить на локальный файл:")
    numbered(doc, "Создайте папку public/images/hosts/ (если её нет)")
    numbered(doc, "Положите файл, например: public/images/hosts/anna.webp")
    numbered(doc, "В db.json замените URL на \"/images/hosts/anna.webp\"")
    para(doc, "Альтернатива: оставить URL Unsplash или указать любой другой URL.")

    h2(doc, "9.3. Популярные направления (главная страница)")
    para(doc, "Файл: src/components/blocks/PopularDestinations/PopularDestinations.tsx")
    para(doc, "Три фото задаются массивом destinationImages с URL Unsplash (строки 6–10).")
    para(doc, "Вариант А — заменить URL на другие фото Unsplash:")
    numbered(doc, "Откройте PopularDestinations.tsx")
    numbered(doc, "Замените URL в массиве destinationImages[0], [1], [2]")
    para(doc, "Вариант Б — использовать локальные файлы:")
    numbered(doc, "Положите файлы в src/assets/images/destinations/ (например spb.webp, moscow.webp, sochi.webp)")
    numbered(doc, "Запустите npm run images:retina")
    numbered(doc, "В PopularDestinations.tsx импортируйте файлы:")
    code_block(doc, "import spbImg from '../../../assets/images/destinations/spb.webp';\n// ...\nconst destinationImages = [spbImg, moscowImg, sochiImg];")

    h2(doc, "9.4. Идеи для путешествий (Travel Ideas)")
    para(doc, "4 тематические статьи с множеством изображений. Файлы лежат в:")
    code_block(doc, "src/assets/images/trip-ideas/\n├── peaks-to-conquer/       # Вершины\n│   ├── card.webp, hero.webp, inline.webp\n│   └── destinations/altai.webp, dombay.webp, kazbek.webp, dolomites.webp\n├── secluded-corners/       # Уединённые уголки\n├── sea-breeze/             # Морской бриз\n└── tropical-paradise/      # Тропический рай")
    para(doc, "Импорты изображений: src/pages/TripIdeaPage/images.ts")
    para(doc, "Данные статей: src/pages/TripIdeaPage/travelIdeas.ts")
    para(doc, "Карточки на главной: src/components/blocks/TravelIdeas/TravelIdeas.tsx (берёт imageUrl из travelIdeas.ts)")

    para(doc, "Пошаговая замена фото идеи «Вершины» (peaks-to-conquer):")
    numbered(doc, "Замените card.webp — миниатюра на главной странице")
    numbered(doc, "Замените hero.webp — большое фото в шапке статьи /ideas/peaks-to-conquer")
    numbered(doc, "Замените inline.webp — иллюстрация внутри текста статьи")
    numbered(doc, "Замените файлы в destinations/ — фото направлений (Алтай, Домбай, Казбек, Дolomites)")
    numbered(doc, "Запустите npm run images:retina для генерации @2x версий")
    numbered(doc, "Перезагрузите страницу в браузере")

    para(doc, "Аналогично для остальных трёх тем: secluded-corners, sea-breeze, tropical-paradise.")
    para(doc, "Важно: не переименовывайте файлы — имена захардкожены в images.ts. Меняйте содержимое файлов на месте.")

    h2(doc, "9.5. Баннер «Сдавайте жильё» (Host Earn Banner)")
    para(doc, "Файлы:")
    code_block(doc, "src/assets/images/home/host-earn.webp\nsrc/assets/images/home/host-earn@2x.webp")
    para(doc, "Используется в: HostEarnBanner.tsx и ListYourSpacePage.tsx")
    numbered(doc, "Замените src/assets/images/home/host-earn.webp")
    numbered(doc, "Запустите npm run images:retina")
    numbered(doc, "Проверьте главную страницу и /list-your-space")

    h2(doc, "9.6. Страницы входа и регистрации")
    para(doc, "LoginPage.tsx и RegisterPage.tsx используют одно фото с Unsplash (боковая панель):")
    code_block(doc, 'src="https://images.unsplash.com/photo-1547448415-e9f5b28e570d?..."')
    para(doc, "Замена на локальный файл:")
    numbered(doc, "Положите фото в src/assets/images/auth/login-hero.webp")
    numbered(doc, "Запустите npm run images:retina")
    numbered(doc, "В LoginPage.tsx и RegisterPage.tsx замените src на импорт:")
    code_block(doc, "import loginHero from '../../assets/images/auth/login-hero.webp';\n// ...\n<img src={loginHero} alt=\"\" />")

    h2(doc, "9.7. Аватар пользователя (профиль)")
    para(doc,
         "Аватар загружается через PersonalInfoTab и хранится в localStorage (userAvatarStorage.ts). "
         "Это не файлы проекта — пользователь загружает фото через интерфейс профиля. "
         "Для демонстрации на защите достаточно загрузить фото в профиле через UI.")

    h2(doc, "9.8. Файл hero.png")
    para(doc,
         "В src/assets/hero.png лежит файл, который в текущей версии кода не используется напрямую. "
         "Можно удалить или заменить, если планируете подключить его в будущем.")

    h2(doc, "9.9. Чек-лист замены всех фото")
    para(doc, "Полный список мест, где нужно заменить изображения:")
    bullet(doc, "26 × card.webp — public/images/properties/{1-26}/")
    bullet(doc, "До 30 × gallery URLs — db.json → images[] (объекты 1–10)")
    bullet(doc, "26 × host.avatar — db.json → host.avatar")
    bullet(doc, "3 × destinationImages — PopularDestinations.tsx")
    bullet(doc, "4 × card + 4 × hero + 4 × inline + 16 × destinations — trip-ideas/")
    bullet(doc, "1 × host-earn.webp — assets/images/home/")
    bullet(doc, "1 × login hero — LoginPage.tsx + RegisterPage.tsx")

    para(doc, "После всех замен:")
    numbered(doc, "npm run images:retina — обновить @2x версии")
    numbered(doc, "npm run dev — проверить в браузере все страницы")
    numbered(doc, "Проверить: главная, /search, /property/1, /ideas/peaks-to-conquer, /login, /list-your-space")

    # 10. Сценарий для защиты
    h1(doc, "10. Сценарий демонстрации на защите")
    h2(doc, "10.1. Что рассказать")
    numbered(doc, "Назначение: SPA для поиска и бронирования жилья, аналог Airbnb")
    numbered(doc, "Стек: React + TypeScript + Vite, Express + PostgreSQL")
    numbered(doc, "Архитектура: компонентный подход, Context API, REST API, i18n")
    numbered(doc, "Данные: 26 объектов в db.json, пользовательские данные в PostgreSQL")

    h2(doc, "10.2. Что показать")
    numbered(doc, "Главная — карусель объектов, популярные направления, идеи путешествий")
    numbered(doc, "Поиск — ввести город, даты, гостей → результаты с фильтрами")
    numbered(doc, "Объект — галерея, календарь с занятыми датами, бронирование")
    numbered(doc, "Регистрация → вход → бронирование → поездки в профиле")
    numbered(doc, "Избранное — добавить/удалить объект")
    numbered(doc, "Mobile — адаптив, мобильный поиск, карусель фото")
    numbered(doc, "Переключение языка ru/en")

    h2(doc, "10.3. Типичные вопросы")
    para(doc, "Где хранятся данные объектов? — src/db.json, 26 записей, загружаются на клиенте.")
    para(doc, "Как работает авторизация? — JWT-токен, bcrypt, PostgreSQL таблица users.")
    para(doc, "Как проверяются занятые даты? — bookedDates в db.json + проверка конфликтов на backend при бронировании.")
    para(doc, "Как устроена адаптивность? — SCSS media queries, breakpoints 767/1023/1024px, отдельные mobile-компоненты.")
    para(doc, "Как заменить фото? — см. раздел 9 данной документации.")

    # Сохранение
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Документ сохранён: {OUTPUT}")


if __name__ == "__main__":
    build()
